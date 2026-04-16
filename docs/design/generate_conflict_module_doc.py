"""Generate the Conflict Intelligence Module design document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style helpers ───────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x47)
BLUE = RGBColor(0x3A, 0x9F, 0xD6)
GREEN = RGBColor(0x27, 0xC4, 0xA0)
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)
MID_GRAY = RGBColor(0x5A, 0x6A, 0x7E)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = NAVY
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, italic=False, color=None, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.runs[0].text = ''
        p.add_run(text)
    return p

def add_quote(text, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = MID_GRAY
    run.font.size = Pt(11)
    if attribution:
        run2 = p.add_run(f'\n\u2014 {attribution}')
        run2.font.color.rgb = MID_GRAY
        run2.font.size = Pt(10)
        run2.bold = True
    return p

def set_cell_shading(cell, color_hex):
    shading = cell._tc.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(s)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, '1B2A47')
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')
    return table

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ARTES')
run.font.size = Pt(14)
run.font.color.rgb = BLUE
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Conflict Intelligence\u2122 Module')
run.font.size = Pt(32)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Theoretical Foundations & Design Document')
run.font.size = Pt(16)
run.font.color.rgb = MID_GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Grounded in the Harvard Negotiation Project\nInterest-Based Negotiation Framework')
run.font.size = Pt(12)
run.font.color.rgb = BLUE
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HeadSoft Tech \u00d7 Helena Coaching')
run.font.size = Pt(11)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('April 2026')
run.font.size = Pt(11)
run.font.color.rgb = MID_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Executive Summary',
    '2. Theoretical Foundations',
    '   2.1 The Harvard Negotiation Project & Interest-Based Negotiation',
    '   2.2 Getting to Yes \u2014 Fisher, Ury & Patton',
    '   2.3 The Three Conversations \u2014 Difficult Conversations (Stone, Patton & Heen)',
    '   2.4 William Ury\u2019s Follow-Up Work',
    '   2.5 Synthesis: From Theory to Software',
    '3. How the Theoretical Framework Maps to the Module',
    '   3.1 Survey Design \u2014 Detecting the Three Conversations',
    '   3.2 AI Analysis \u2014 Interest-Based Conflict Diagnosis',
    '   3.3 Manager Conversation Guide \u2014 Principled Negotiation in Practice',
    '   3.4 Recommended Actions \u2014 From Positions to Interests',
    '   3.5 Escalation Pathway \u2014 The Third Side',
    '4. Module Architecture',
    '   4.1 Data Flow',
    '   4.2 AI Prompt Design Philosophy',
    '   4.3 Risk Scoring Model',
    '5. Survey Instruments',
    '   5.1 Bi-Weekly Pulse Survey (15 items)',
    '   5.2 Quarterly Deep-Dive Analysis (21 items)',
    '   5.3 Question Design Rationale',
    '6. The Escalation & Mediation Pathway',
    '7. Skill Building & Development',
    '8. Ethical Guardrails & Privacy',
    '9. References & Further Reading',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith('   '):
        run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'The ARTES Conflict Intelligence\u2122 module is a prediction-and-action system that treats '
    'workplace conflict as a predictable, measurable, and ultimately preventable organisational '
    'dynamic \u2014 not a failure of management. It combines anonymous survey data, AI-driven '
    'pattern recognition, and structured intervention frameworks to help organisations move from '
    'reactive firefighting to proactive conflict stewardship.'
)

doc.add_paragraph(
    'The module\u2019s theoretical backbone is the interest-based negotiation model pioneered by '
    'the Harvard Negotiation Project, most widely known through Fisher, Ury and Patton\u2019s '
    'Getting to Yes (1981, revised 2011). Every layer of the system \u2014 from survey question '
    'design to AI prompt engineering to manager conversation guides \u2014 is built on the '
    'principle that durable conflict resolution requires understanding underlying interests, '
    'not merely managing surface-level positions.'
)

doc.add_paragraph(
    'This document traces the line from foundational theory to working software: how the '
    'Harvard Negotiation Project\u2019s research, the Three Conversations framework from '
    'Difficult Conversations (Stone, Patton & Heen), and William Ury\u2019s broader body of '
    'work on negotiation and mediation inform every design decision in the module.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. THEORETICAL FOUNDATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Theoretical Foundations', level=1)

# 2.1
doc.add_heading('2.1 The Harvard Negotiation Project & Interest-Based Negotiation', level=2)

doc.add_paragraph(
    'The Harvard Negotiation Project (HNP), founded at Harvard Law School in 1979, is the '
    'seminal academic initiative behind interest-based (or principled) negotiation. Its central '
    'insight \u2014 that effective negotiation separates people from problems, focuses on '
    'interests rather than positions, generates options for mutual gain, and insists on '
    'objective criteria \u2014 has become the foundation of modern conflict resolution practice '
    'across legal, diplomatic, organisational, and interpersonal domains.'
)

doc.add_paragraph(
    'Interest-based negotiation stands in direct contrast to positional bargaining, where '
    'each party stakes out a position and makes concessions. Positional bargaining tends to '
    'produce suboptimal outcomes, damage relationships, and entrench adversarial dynamics. '
    'The interest-based model instead asks: What does each party actually need? What are the '
    'underlying concerns, fears, desires, and constraints that drive the stated positions? '
    'By reframing conflict around interests, parties can often discover solutions that '
    'positional bargaining would never surface.'
)

add_para(
    'The four principles of the Harvard method form the architectural spine of '
    'the Conflict Intelligence module:',
    bold=True, color=NAVY, size=Pt(11), space_after=4
)

add_bullet('Separate the people from the problem.', bold_prefix='Principle 1: ')
doc.add_paragraph(
    'Emotions, perceptions, and communication breakdowns are real obstacles to resolution, but '
    'they are not the substance of the dispute. The module\u2019s survey instruments measure '
    'psychological safety, trust, and communication quality as distinct dimensions from the '
    'structural or procedural issues that generate conflict. This separation allows the AI '
    'analysis to diagnose both relational and structural dynamics independently.'
).paragraph_format.left_indent = Cm(1.5)

add_bullet('Focus on interests, not positions.', bold_prefix='Principle 2: ')
doc.add_paragraph(
    'Survey questions are deliberately designed to surface underlying needs and concerns rather '
    'than asking employees to identify a \u201cwinner\u201d or \u201cblame party.\u201d The AI '
    'prompts instruct Claude to identify root causes and unmet needs behind conflict patterns, '
    'not to adjudicate who is right. Manager conversation guides are framed as curiosity-driven '
    'dialogues that explore interests, not interrogations that assign fault.'
).paragraph_format.left_indent = Cm(1.5)

add_bullet('Generate options for mutual gain.', bold_prefix='Principle 3: ')
doc.add_paragraph(
    'The Recommended Actions engine is structured to produce creative, multi-stakeholder '
    'solutions across different time horizons. Actions are assigned to roles (HR, Manager, '
    'Coach, Team Lead) and prioritised by urgency, ensuring that resolution is a shared '
    'responsibility rather than a single person\u2019s burden. Preventive measures address '
    'systemic conditions that generate conflict, not just individual incidents.'
).paragraph_format.left_indent = Cm(1.5)

add_bullet('Insist on objective criteria.', bold_prefix='Principle 4: ')
doc.add_paragraph(
    'The risk scoring model (0\u2013100) provides an objective, data-driven baseline that '
    'removes subjectivity from the initial assessment. Aggregated survey data serves as the '
    'objective input; the AI analysis layer interprets patterns against established conflict '
    'research. Follow-up pulse surveys at 30\u201360 days measure whether interventions have '
    'moved the score, creating a feedback loop grounded in evidence rather than opinion.'
).paragraph_format.left_indent = Cm(1.5)

# 2.2
doc.add_heading('2.2 Getting to Yes \u2014 Fisher, Ury & Patton', level=2)

add_quote(
    'The most powerful interests are basic human needs. In searching for the basic interests '
    'behind a declared position, look particularly for those bedrock concerns which motivate '
    'all people: security, economic well-being, a sense of belonging, recognition, and control '
    'over one\u2019s life.',
    'Fisher, Ury & Patton, Getting to Yes (2011)'
)

doc.add_paragraph(
    'Getting to Yes (first published 1981, revised editions 1991 and 2011) is the most widely '
    'read negotiation text in history, translated into 36 languages and used as a foundational '
    'resource in law schools, business schools, and diplomatic training programmes worldwide. '
    'Its enduring contribution is the codification of principled negotiation into a teachable, '
    'repeatable method.'
)

doc.add_paragraph(
    'The book\u2019s core framework maps directly onto workplace conflict resolution:'
)

add_table(
    ['Getting to Yes Concept', 'Workplace Application', 'Module Implementation'],
    [
        ['Separate People from Problem', 'Address emotions and miscommunication independently from structural issues',
         'Survey dimensions measure psychological safety, communication, and structural stressors as separate constructs'],
        ['Focus on Interests', 'Ask why employees feel the way they do, not just what they want',
         'AI narrative identifies root causes and unmet needs; manager scripts use open-ended interest-probing questions'],
        ['Invent Options for Mutual Gain', 'Brainstorm solutions that address multiple parties\u2019 core concerns',
         'Recommended Actions engine generates multi-stakeholder, time-phased interventions across roles'],
        ['Use Objective Criteria', 'Ground discussion in data, not power or personality',
         'Risk score (0\u2013100), aggregated survey data, and follow-up measurement provide evidence-based anchors'],
        ['BATNA (Best Alternative)', 'Understand what happens if conflict is not resolved',
         'Escalation pathway provides a structured alternative; risk level (high/critical) signals when BATNA is active'],
    ]
)

doc.add_paragraph()

doc.add_paragraph(
    'Fisher and Ury\u2019s emphasis on \u201cinventing options before deciding\u201d is '
    'particularly important for the module\u2019s design. The AI does not recommend a single '
    'solution; it generates a portfolio of actions across time horizons, allowing managers and '
    'HR to select, adapt, and combine interventions based on their contextual knowledge. This '
    'mirrors the Getting to Yes principle that good negotiation expands the pie before '
    'dividing it.'
)

# 2.3
doc.add_heading('2.3 The Three Conversations \u2014 Difficult Conversations (Stone, Patton & Heen)', level=2)

add_quote(
    'Difficult conversations are almost never about getting the facts right. They are about '
    'conflicting perceptions, interpretations, and values. They are not about what a contract '
    'says, they are about what a contract means.',
    'Stone, Patton & Heen, Difficult Conversations (1999)'
)

doc.add_paragraph(
    'Difficult Conversations (1999), authored by Douglas Stone, Bruce Patton, and Sheila Heen '
    '\u2014 all members of the Harvard Negotiation Project \u2014 extends the interest-based '
    'model into the domain of interpersonal workplace conflict. Its central contribution is '
    'the Three Conversations framework: every difficult conversation is actually three '
    'simultaneous conversations operating beneath the surface.'
)

add_para('The Three Conversations:', bold=True, color=NAVY, space_after=4)

p = doc.add_paragraph()
run = p.add_run('1. The \u201cWhat Happened?\u201d Conversation')
run.bold = True
run.font.color.rgb = NAVY
doc.add_paragraph(
    'Each party has a different story about what happened, who intended what, and who is to '
    'blame. The core error is the assumption that the other party\u2019s intentions are obvious '
    'from their actions. In workplace conflict, this manifests as attribution errors: '
    '\u201cThey didn\u2019t respond to my email because they don\u2019t respect my work\u201d '
    'versus the reality that they were overwhelmed with competing deadlines.'
)
doc.add_paragraph(
    'Module implementation: The survey instruments do not ask employees to narrate events or '
    'assign blame. Instead, they measure behavioural indicators (e.g., \u201cDisagreements in '
    'my team are resolved constructively\u201d) that surface the \u201cWhat Happened\u201d '
    'conversation without requiring parties to agree on a shared narrative. The AI analysis '
    'identifies patterns across multiple respondents, providing a composite picture that no '
    'single individual\u2019s story could provide.'
)

p = doc.add_paragraph()
run = p.add_run('2. The Feelings Conversation')
run.bold = True
run.font.color.rgb = NAVY
doc.add_paragraph(
    'Emotions are not a side effect of conflict; they are the substance of it. Unacknowledged '
    'feelings drive escalation, withdrawal, and retaliation. In the workplace, the Feelings '
    'Conversation is often suppressed by professional norms (\u201cdon\u2019t be emotional\u201d), '
    'which means it operates underground \u2014 manifesting as disengagement, passive aggression, '
    'or sudden escalation when feelings finally break through.'
)
doc.add_paragraph(
    'Module implementation: The Psychological Safety dimension of the survey directly measures '
    'the conditions under which feelings can be expressed (\u201cI feel safe expressing concerns '
    'or disagreements without fear of negative consequences\u201d). The Wellbeing dimension '
    'captures the emotional impact of unresolved conflict. Together, these dimensions give the '
    'AI the signal it needs to identify suppressed Feelings Conversations. Manager scripts '
    'explicitly include emotion-acknowledging language: \u201cI can see this has been '
    'frustrating\u2014can you help me understand what\u2019s been most difficult for you?\u201d'
)

p = doc.add_paragraph()
run = p.add_run('3. The Identity Conversation')
run.bold = True
run.font.color.rgb = NAVY
doc.add_paragraph(
    'The most hidden and powerful layer. Every difficult conversation threatens something about '
    'how we see ourselves: Am I competent? Am I a good person? Am I worthy of respect? When '
    'identity is at stake, people become defensive, rigid, or avoidant \u2014 not because the '
    'substantive issue is complex, but because their self-image is under threat.'
)
doc.add_paragraph(
    'Module implementation: Survey questions about belonging, recognition, and value (\u201cI '
    'believe my unique perspective and contributions are genuinely valued by my team\u201d) '
    'tap directly into the Identity Conversation. The module\u2019s escalation pathway and '
    'coaching integration recognise that identity-level conflicts require skilled facilitation, '
    'not just procedural resolution. When the AI detects patterns suggesting identity threat '
    '(low belonging scores combined with high avoidance), it flags the need for deeper '
    'intervention.'
)

add_table(
    ['Three Conversations', 'What It Sounds Like at Work', 'Survey Dimensions That Detect It', 'AI Response'],
    [
        ['What Happened?', '\u201cThat\u2019s not what was agreed\u201d / \u201cYou never told me\u201d',
         'Communication & Trust, Conflict Frequency',
         'Identifies divergent narratives without assigning blame; suggests structured dialogue'],
        ['Feelings', '\u201cI don\u2019t feel heard\u201d / withdrawal / sudden outbursts',
         'Psychological Safety, Wellbeing',
         'Flags suppressed emotional dynamics; includes emotion-acknowledging language in scripts'],
        ['Identity', '\u201cThey don\u2019t respect my expertise\u201d / defensiveness / rigidity',
         'Wellbeing & Belonging, Interpersonal Dynamics',
         'Detects identity threat patterns; escalates to coaching/mediation when present'],
    ]
)

doc.add_paragraph()

# 2.4
doc.add_heading('2.4 William Ury\u2019s Follow-Up Work', level=2)

doc.add_paragraph(
    'William Ury\u2019s contributions extend well beyond Getting to Yes. His subsequent works '
    'provide the conceptual tools for several key features of the Conflict Intelligence module:'
)

p = doc.add_paragraph()
run = p.add_run('Getting Past No (1991)')
run.bold = True
doc.add_paragraph(
    'Addresses what to do when the other party won\u2019t cooperate \u2014 the practical '
    'reality of most workplace conflict. Ury\u2019s five-step \u201cbreakthrough negotiation\u201d '
    'strategy (Go to the Balcony, Step to Their Side, Reframe, Build Them a Golden Bridge, '
    'Use Power to Educate) directly informs the module\u2019s manager conversation scripts. '
    'The \u201cGo to the Balcony\u201d concept \u2014 stepping back from reactive emotions to '
    'gain perspective \u2014 is embedded in the module\u2019s design: the AI analysis serves as '
    'a \u201cbalcony view\u201d of the conflict, giving managers an objective vantage point '
    'before they enter the conversation.'
)

p = doc.add_paragraph()
run = p.add_run('The Third Side (2000)')
run.bold = True
doc.add_paragraph(
    'Ury argues that conflict is not just between two parties; the surrounding community is '
    'always the \u201cthird side\u201d that can contain, prevent, or resolve disputes. In '
    'organisational terms, the third side includes HR, coaches, peers, and organisational '
    'culture itself. The module\u2019s escalation pathway is a direct implementation of the '
    'Third Side model: when bilateral resolution fails, the system activates structured '
    'third-party intervention (Helena\u2019s mediation) rather than leaving parties to escalate '
    'unilaterally. The recommended actions are assigned across multiple roles (Manager, HR, '
    'Coach, Team Lead) precisely because Ury\u2019s research shows that effective conflict '
    'resolution is a community function, not a bilateral transaction.'
)

p = doc.add_paragraph()
run = p.add_run('Getting to Yes with Yourself (2015)')
run.bold = True
doc.add_paragraph(
    'Ury\u2019s most recent major work addresses the inner negotiation that precedes the outer '
    'one: before you can negotiate effectively with others, you must negotiate with yourself. '
    'This maps directly to the Identity Conversation from Difficult Conversations and informs '
    'the module\u2019s coaching integration. The Skill Building section includes self-reflective '
    'exercises drawn from this framework, and the coaching pathway (via Helena) provides '
    'the facilitated self-awareness work that Ury identifies as the prerequisite for effective '
    'conflict engagement.'
)

# 2.5
doc.add_heading('2.5 Synthesis: From Theory to Software', level=2)

doc.add_paragraph(
    'The Conflict Intelligence module is not a digital textbook. It operationalises these '
    'theoretical frameworks into a continuous feedback system:'
)

add_bullet(
    'Survey instruments detect the Three Conversations at scale, across teams and time periods, '
    'without requiring anyone to narrate their side of the story.'
)
add_bullet(
    'AI analysis applies interest-based reasoning to aggregated data, identifying the '
    'underlying interests and structural conditions that generate conflict \u2014 not just the '
    'visible symptoms.'
)
add_bullet(
    'Manager conversation guides translate principled negotiation into practical, repeatable '
    'dialogue scripts that any manager can use, regardless of their training.'
)
add_bullet(
    'Recommended actions implement the \u201cinvent options for mutual gain\u201d principle at '
    'an organisational level, distributing responsibility across roles and time horizons.'
)
add_bullet(
    'The escalation pathway activates the Third Side when bilateral resolution is insufficient.'
)
add_bullet(
    'Follow-up measurement closes the loop with objective criteria, measuring whether '
    'interventions have moved the risk score.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. HOW THE THEORETICAL FRAMEWORK MAPS TO THE MODULE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. How the Theoretical Framework Maps to the Module', level=1)

# 3.1
doc.add_heading('3.1 Survey Design \u2014 Detecting the Three Conversations', level=2)

doc.add_paragraph(
    'The module\u2019s two survey instruments (Bi-Weekly Pulse and Quarterly Deep-Dive) are '
    'structured to detect all three conversation layers identified by Stone, Patton and Heen, '
    'using behavioural indicators rather than narrative accounts.'
)

add_table(
    ['Survey Category', 'Conversation Layer', 'What It Measures', 'Key Questions'],
    [
        ['Psychological Safety', 'Feelings + Identity',
         'Whether employees can express concerns without fear; whether identity feels safe',
         'cp01, cp02, cd05, cd06, cd07'],
        ['Communication & Trust', 'What Happened?',
         'Whether information flows accurately; whether commitments are honoured',
         'cp04, cp05, cp06'],
        ['Conflict Frequency', 'What Happened? (intensity)',
         'How often conflict surfaces; whether it\u2019s resolved or accumulating',
         'cp07, cp08, cp09'],
        ['Management Effectiveness', 'Third Side capacity',
         'Whether managers model and facilitate resolution',
         'cp10, cp11, cd11, cd12'],
        ['Escalation Intent', 'All three (threshold indicator)',
         'Whether someone has reached the point of needing external intervention',
         'cp12, cp13'],
        ['Wellbeing & Belonging', 'Feelings + Identity',
         'Emotional impact of unresolved conflict; sense of inclusion',
         'cp14, cp15, cd20, cd21'],
        ['Interpersonal Dynamics', 'Identity + What Happened?',
         'Power imbalances, persistent tensions, collaboration under disagreement',
         'cd08, cd09, cd10'],
        ['Workload & Structural Stressors', 'What Happened? (structural)',
         'Whether conflict is generated by unclear roles, competing priorities, or resource constraints',
         'cd15, cd16, cd17'],
    ]
)

doc.add_paragraph()

doc.add_paragraph(
    'This design ensures that the AI receives signals from all three conversation layers, '
    'enabling it to produce analyses that address the full depth of the conflict \u2014 not '
    'just the surface-level \u201cWhat Happened\u201d story that most conflict systems capture.'
)

# 3.2
doc.add_heading('3.2 AI Analysis \u2014 Interest-Based Conflict Diagnosis', level=2)

doc.add_paragraph(
    'The AI prompts are explicitly engineered to produce interest-based analysis. The system '
    'prompt instructs Claude to use \u201cHelena\u2019s coaching-integrated, interest-based '
    'mediation methodology\u201d and to identify root causes and underlying needs rather than '
    'assigning blame or adjudicating disputes.'
)

doc.add_paragraph(
    'The analysis output is structured to separate the three analytical layers:'
)

add_bullet(
    'The AI Narrative (2\u20133 paragraphs) provides a composite interpretation of the '
    'conflict dynamics, written in the language of interests and needs rather than positions '
    'and blame. It describes what the data suggests about the team\u2019s underlying dynamics, '
    'not what any individual did wrong.',
    bold_prefix='Narrative layer: '
)
add_bullet(
    'The Risk Score (0\u2013100) and Risk Level provide the objective criteria that Fisher, '
    'Ury and Patton insist upon \u2014 a data-grounded starting point for conversation that '
    'removes the subjectivity of \u201cI think things are fine\u201d vs. \u201cI think '
    'things are terrible.\u201d',
    bold_prefix='Quantitative layer: '
)
add_bullet(
    'The Conflict Types array identifies specific patterns (e.g., \u201cRole Ambiguity,\u201d '
    '\u201cCommunication Breakdown,\u201d \u201cLeadership-Process Gap\u201d) that name the '
    'structural interests at stake without personalising the conflict.',
    bold_prefix='Pattern layer: '
)

# 3.3
doc.add_heading('3.3 Manager Conversation Guide \u2014 Principled Negotiation in Practice', level=2)

doc.add_paragraph(
    'The Manager Script output translates the interest-based framework into practical '
    'conversation guidance. For sub-analyses, the script is structured as:'
)

add_bullet(
    'How to open the conversation \u2014 framing the discussion around shared concerns '
    'rather than accusations (Principle 1: Separate People from Problem)',
    bold_prefix='Opening: '
)
add_bullet(
    'Interest-based questions that explore underlying needs: \u201cWhat would need to '
    'change for you to feel confident about role boundaries?\u201d rather than \u201cWho '
    'isn\u2019t doing their job?\u201d (Principle 2: Focus on Interests)',
    bold_prefix='Key Questions: '
)
add_bullet(
    'Specific resolution approaches and next steps that invite collaborative problem-solving '
    '(Principle 3: Invent Options for Mutual Gain)',
    bold_prefix='Resolution: '
)

doc.add_paragraph(
    'This structure means that even a manager with no negotiation training can conduct a '
    'conversation that follows the Harvard model\u2019s core principles. The script is a '
    'scaffold, not a screenplay \u2014 it provides direction while leaving space for the '
    'manager\u2019s contextual judgment and the employee\u2019s authentic response.'
)

# 3.4
doc.add_heading('3.4 Recommended Actions \u2014 From Positions to Interests', level=2)

doc.add_paragraph(
    'The three-tier action structure (Immediate, Short-Term, Long-Term) plus Preventive '
    'Measures is designed to address conflict at every level of the interest hierarchy:'
)

add_table(
    ['Time Horizon', 'Interest Level Addressed', 'Example', 'Harvard Principle'],
    [
        ['Immediate (This Week)', 'Acute emotional/safety needs',
         'Schedule 1:1 check-ins with each team member to understand their perspective',
         'Separate People from Problem'],
        ['Short-Term (2\u20134 Weeks)', 'Structural and procedural interests',
         'Clarify role boundaries via RACI matrix co-created with the team',
         'Focus on Interests + Invent Options'],
        ['Long-Term (1\u20133 Months)', 'Cultural and systemic interests',
         'Implement quarterly conflict health surveys to track trends over time',
         'Objective Criteria'],
        ['Preventive Measures', 'Systemic conditions that generate conflict',
         'Establish a team agreement for how disagreements will be raised and resolved',
         'All four principles (preventive)'],
    ]
)

doc.add_paragraph()

# 3.5
doc.add_heading('3.5 Escalation Pathway \u2014 The Third Side', level=2)

doc.add_paragraph(
    'William Ury\u2019s Third Side framework recognises that some conflicts cannot be resolved '
    'bilaterally. The module\u2019s escalation pathway activates when the risk level is high or '
    'critical, providing a structured handoff to professional mediation:'
)

add_table(
    ['Step', 'Action', 'Theoretical Basis'],
    [
        ['1. Flag for Escalation', 'HR/Manager clicks escalate; system records request',
         'Recognition that bilateral resolution has reached its limits (Ury\u2019s \u201cwhen to bring in the Third Side\u201d)'],
        ['2. Initial Consultation', 'Helena reviews AI analysis; 30-minute intake with HR',
         'Third Side enters with data-grounded understanding (Objective Criteria principle)'],
        ['3. Mediation Process', 'Interest-based mediation session facilitated by Helena',
         'Full Harvard method: separate people from problem, explore interests, generate options'],
        ['4. Resolution & Follow-Up', 'Agreement documented; pulse survey at 30\u201360 days',
         'Objective Criteria: measure whether interests have been addressed, not just positions settled'],
    ]
)

doc.add_paragraph()

doc.add_paragraph(
    'The escalation pathway is deliberately not automated. Conflict at the high/critical level '
    'involves identity-level dynamics (the Identity Conversation) that require human judgment, '
    'empathy, and relational skill. The AI provides the diagnostic intelligence; a skilled '
    'human provides the facilitation.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. MODULE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Module Architecture', level=1)

doc.add_heading('4.1 Data Flow', level=2)

doc.add_paragraph(
    'The module follows a five-stage pipeline that mirrors the interest-based '
    'negotiation process itself:'
)

stages = [
    ('Intake', 'Anonymous surveys distributed to teams (min. 5 respondents for privacy). '
     'No individual data is exposed. This parallels the \u201cactive listening\u201d phase of '
     'interest-based negotiation: gathering all perspectives without judgment.'),
    ('Aggregation', 'Responses are averaged per question, producing a numerical profile of '
     'the team\u2019s conflict dynamics. Individual responses are never surfaced. This creates '
     'the \u201cobjective criteria\u201d that Fisher and Ury require for principled negotiation.'),
    ('AI Analysis', 'Claude analyses the aggregated data against the interest-based framework, '
     'producing a risk score (0\u2013100), risk level, detected conflict types, a narrative, '
     'and a manager conversation guide.'),
    ('Action', 'The Recommended Actions engine generates prioritised, multi-stakeholder '
     'interventions. Managers and HR select, adapt, and execute. Completion is tracked via '
     'per-action checkboxes.'),
    ('Follow-Up', 'Subsequent pulse surveys measure whether interventions have moved the '
     'risk score, closing the feedback loop with evidence rather than opinion.'),
]
for title, desc in stages:
    p = doc.add_paragraph()
    run = p.add_run(f'Stage {stages.index((title, desc)) + 1}: {title}. ')
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(desc)

doc.add_heading('4.2 AI Prompt Design Philosophy', level=2)

doc.add_paragraph(
    'Every AI prompt in the module is engineered around three principles drawn from '
    'the theoretical foundations:'
)

add_bullet(
    'Prompts never ask the AI to determine who is right or wrong. They ask it to identify '
    'patterns, root causes, and unmet needs.',
    bold_prefix='Interest-based framing: '
)
add_bullet(
    'Prompts instruct the AI to produce manager scripts using open-ended, interest-probing '
    'questions rather than accusatory or position-defending language.',
    bold_prefix='Conversation-ready output: '
)
add_bullet(
    'Prompts require the AI to generate multiple options across time horizons and stakeholder '
    'roles, reflecting the \u201cinvent options for mutual gain\u201d principle.',
    bold_prefix='Multi-option generation: '
)

doc.add_heading('4.3 Risk Scoring Model', level=2)

doc.add_paragraph(
    'The risk score (0\u2013100) is a composite measure that the AI derives from the '
    'aggregated survey data. The four risk levels map to escalating response protocols:'
)

add_table(
    ['Risk Level', 'Score Range', 'Interpretation', 'Recommended Response'],
    [
        ['Low', '0\u201330', 'Healthy conflict culture; minor tensions being managed',
         'Monitor with regular pulse surveys; reinforce positive practices'],
        ['Medium', '31\u201355', 'Emerging patterns that could escalate without attention',
         'Manager-led conversations using provided scripts; short-term actions'],
        ['High', '56\u201375', 'Active conflict affecting productivity and wellbeing',
         'Immediate actions + coaching involvement; consider escalation to mediation'],
        ['Critical', '76\u2013100', 'Severe conflict requiring urgent intervention',
         'Escalate to professional mediation; immediate safety and wellbeing measures'],
    ]
)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. SURVEY INSTRUMENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Survey Instruments', level=1)

doc.add_heading('5.1 Bi-Weekly Pulse Survey (15 items)', level=2)

doc.add_paragraph(
    'The pulse survey is designed for frequent, low-friction measurement. It takes 3\u20135 '
    'minutes to complete and covers six dimensions:'
)

pulse_cats = [
    ('Psychological Safety', 3, 'cp01\u2013cp03',
     'Can employees express concerns without fear? Have they witnessed unsafe behaviour?'),
    ('Communication & Trust', 3, 'cp04\u2013cp06',
     'Is communication open and honest? Are commitments honoured? Are disagreements resolved constructively?'),
    ('Conflict Frequency', 3, 'cp07\u2013cp09',
     'How often does tension occur? Is there unresolved conflict? Is productivity affected?'),
    ('Management Effectiveness', 2, 'cp10\u2013cp11',
     'Does the manager address conflict fairly? Does the manager create a respectful environment?'),
    ('Escalation Intent', 2, 'cp12\u2013cp13',
     'Does the situation need HR/leadership intervention? What type of conflict is it?'),
    ('Wellbeing', 2, 'cp14\u2013cp15',
     'Overall wellbeing and belonging rating; open-ended comments (anonymous)'),
]

add_table(
    ['Category', 'Items', 'Question IDs', 'What It Measures'],
    [[c, str(n), ids, m] for c, n, ids, m in pulse_cats]
)

doc.add_paragraph()

doc.add_heading('5.2 Quarterly Deep-Dive Analysis (21 items)', level=2)

doc.add_paragraph(
    'The quarterly survey provides deeper structural analysis across seven dimensions, '
    'designed to surface systemic patterns that bi-weekly pulses may miss:'
)

deep_cats = [
    ('Conflict Culture', 4, 'cd01\u2013cd04',
     'Organisational norms around conflict; training availability'),
    ('Psychological Safety', 3, 'cd05\u2013cd07',
     'Safety to challenge seniors; learning from mistakes; feeling valued'),
    ('Interpersonal Dynamics', 3, 'cd08\u2013cd10',
     'Persistent tensions; power imbalances; ability to disagree productively'),
    ('Leadership & Mediation', 4, 'cd11\u2013cd14',
     'Manager modelling; handling of complaints; trust in HR; avoidance of escalation'),
    ('Workload & Structural Stressors', 3, 'cd15\u2013cd17',
     'Role ambiguity; competing priorities; resource adequacy'),
    ('Cross-Team Conflict', 2, 'cd18\u2013cd19',
     'Inter-departmental friction; cross-team conflict experience'),
    ('Outcomes & Impact', 2, 'cd20\u2013cd21',
     'Impact on engagement and retention; single most important change (open-ended)'),
]

add_table(
    ['Category', 'Items', 'Question IDs', 'What It Measures'],
    [[c, str(n), ids, m] for c, n, ids, m in deep_cats]
)

doc.add_paragraph()

doc.add_heading('5.3 Question Design Rationale', level=2)

doc.add_paragraph(
    'Every survey question was designed against the interest-based framework. Key principles:'
)

add_bullet(
    'Questions measure behavioural indicators, not attribution. \u201cDisagreements in my team '
    'are resolved constructively\u201d is observable and aggregatable; \u201cWho causes the most '
    'conflict?\u201d is not.',
    bold_prefix='Behavioural, not attributional: '
)
add_bullet(
    'Questions measure across all three conversation layers (What Happened, Feelings, Identity) '
    'to give the AI a complete signal.',
    bold_prefix='Multi-layer coverage: '
)
add_bullet(
    'Minimum 5 respondents per analysis. Questions are designed so that no individual response '
    'can be identified from the aggregate.',
    bold_prefix='Anonymity by design: '
)
add_bullet(
    'Both instruments include boolean and open-text questions alongside scales, allowing the '
    'detection of qualitative signals that pure numeric scales miss.',
    bold_prefix='Mixed methods: '
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. ESCALATION & MEDIATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. The Escalation & Mediation Pathway', level=1)

doc.add_paragraph(
    'The escalation pathway is the module\u2019s implementation of William Ury\u2019s Third '
    'Side principle. It activates when the conflict has exceeded the capacity of the immediate '
    'parties to resolve it \u2014 typically at the high or critical risk level.'
)

doc.add_paragraph(
    'The pathway is deliberately human-centred. AI provides the diagnostic intelligence, but '
    'the mediation itself is conducted by a skilled professional (Helena) using the full '
    'interest-based mediation methodology:'
)

add_bullet(
    'Separate sessions with each party to understand their interests, fears, and needs',
    bold_prefix='Pre-mediation: '
)
add_bullet(
    'Joint session structured around interest exploration (not position defence), option '
    'generation, and criteria-based evaluation of solutions',
    bold_prefix='Mediation: '
)
add_bullet(
    'Written agreement that addresses underlying interests, with measurable commitments '
    'and a follow-up timeline',
    bold_prefix='Agreement: '
)
add_bullet(
    'Pulse survey at 30\u201360 days to measure whether the risk score has moved, '
    'providing objective evidence of resolution',
    bold_prefix='Follow-up: '
)

doc.add_paragraph(
    'The system tracks escalation status through four stages: Pending \u2192 In Progress '
    '\u2192 Resolved/Escalated. This provides organisational visibility into the conflict '
    'resolution pipeline without exposing the substance of the mediation.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. SKILL BUILDING
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Skill Building & Development', level=1)

doc.add_paragraph(
    'The module includes a Skill Building section that provides resources grounded in '
    'the interest-based framework. These resources are designed to build organisational '
    'capacity for conflict resolution over time, reducing dependence on escalation.'
)

doc.add_heading('Interest-Based Negotiation Toolkit', level=3)
add_bullet('Interest Mapping Worksheet \u2014 helps parties identify their own and the other side\u2019s underlying interests')
add_bullet('BATNA Assessment Guide \u2014 structured evaluation of alternatives to agreement')
add_bullet('Reframing Exercises \u2014 practice converting positional statements into interest-based questions')
add_bullet('Manager Conversation Templates \u2014 pre-structured scripts for common conflict scenarios')
add_bullet('The Balcony Technique \u2014 guided exercise for gaining emotional distance before engaging')

doc.add_heading('Coaching Integration', level=3)
doc.add_paragraph(
    'The module integrates with the broader ARTES coaching platform. Coaches can generate '
    'conflict-focused Individual Development Plans (IDPs) using the GROW model, drawing on '
    'analysis data to set evidence-based development goals around conflict competence.'
)

doc.add_heading('External Assessment Integration', level=3)
doc.add_paragraph(
    'The module supports integration with established conflict and emotional intelligence '
    'assessments:'
)
add_bullet('Thomas-Kilmann Instrument (TKI) \u2014 conflict handling mode assessment')
add_bullet('EQ-i 2.0 \u2014 emotional intelligence assessment with conflict-relevant subscales')
add_bullet('DISC \u2014 behavioural style assessment for communication adaptation')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. ETHICAL GUARDRAILS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Ethical Guardrails & Privacy', level=1)

doc.add_paragraph(
    'The module embeds ethical safeguards at every layer, consistent with the interest-based '
    'principle that fair process is not optional \u2014 it is a prerequisite for durable resolution.'
)

add_bullet(
    'All survey responses are anonymous. No individual response is ever surfaced in analysis. '
    'Minimum group size of 5 respondents per analysis prevents statistical de-anonymisation.',
    bold_prefix='Anonymity: '
)
add_bullet(
    'The AI is a diagnostic tool, not an adjudicator. It identifies patterns and suggests '
    'frameworks; it does not determine fault, recommend discipline, or make decisions about '
    'individuals. All AI outputs explicitly state that they are starting points for human '
    'judgment, not conclusions.',
    bold_prefix='AI as advisor, not judge: '
)
add_bullet(
    'Analysis data is tenant-isolated. No cross-organisation data sharing occurs. Each '
    'organisation\u2019s data is stored and processed independently.',
    bold_prefix='Data isolation: '
)
add_bullet(
    'Escalation is always voluntary. The system flags high-risk situations but never '
    'auto-escalates. The decision to involve a mediator or coach rests with the human '
    'stakeholders.',
    bold_prefix='Human agency: '
)
add_bullet(
    'Manager scripts are scaffolds, not mandates. They provide suggested language and '
    'structure, but explicitly preserve space for contextual adaptation. The goal is to '
    'equip managers, not to script them.',
    bold_prefix='Manager autonomy: '
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. References & Further Reading', level=1)

doc.add_heading('Foundational Works', level=3)

refs = [
    'Fisher, R., Ury, W. & Patton, B. (2011). Getting to Yes: Negotiating Agreement Without Giving In. 3rd edition. Penguin Books.',
    'Stone, D., Patton, B. & Heen, S. (1999). Difficult Conversations: How to Discuss What Matters Most. Viking Penguin.',
    'Ury, W. (1991). Getting Past No: Negotiating in Difficult Situations. Bantam Books.',
    'Ury, W. (2000). The Third Side: Why We Fight and How We Can Stop. Penguin Books.',
    'Ury, W. (2015). Getting to Yes with Yourself: And Other Worthy Opponents. HarperOne.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)

doc.add_heading('Supporting Literature', level=3)

supp_refs = [
    'Heen, S. & Stone, D. (2014). Thanks for the Feedback: The Science and Art of Receiving Feedback Well. Viking.',
    'Mnookin, R., Peppet, S. & Tulumello, A. (2000). Beyond Winning: Negotiating to Create Value in Deals and Disputes. Harvard University Press.',
    'Shapiro, D. (2016). Negotiating the Nonnegotiable: How to Resolve Your Most Emotionally Charged Conflicts. Viking.',
    'Edmondson, A. (2018). The Fearless Organization: Creating Psychological Safety in the Workplace. Wiley.',
    'Thomas, K.W. & Kilmann, R.H. (1974). Thomas-Kilmann Conflict Mode Instrument. CPP, Inc.',
]
for ref in supp_refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 End of Document \u2014')
run.font.color.rgb = MID_GRAY
run.italic = True

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r'D:\CUSTOMERS\HeadSoft\people-intelligence-platform\docs\design\conflict_intelligence_module.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
