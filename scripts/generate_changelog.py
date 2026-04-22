#!/usr/bin/env python3
"""Generate a weekly changelog .docx for the ARTES platform.

Usage:
    python3 scripts/generate_changelog.py [--days N] [--since YYYY-MM-DD] [--until YYYY-MM-DD]

Defaults to the past 7 days. Output is saved to docs/changelog_YYYY-MM-DD.docx.
Requires: pip install python-docx
"""

import argparse
import re
import subprocess
from collections import OrderedDict
from datetime import date, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

NAVY  = RGBColor(0x1B, 0x2A, 0x47)
BLUE  = RGBColor(0x3A, 0x9F, 0xD6)
GREY  = RGBColor(0x6B, 0x7C, 0x93)
LGREY = RGBColor(0x9A, 0xA5, 0xB4)

# Ordered sections with keywords matched against subject + scope + body.
# Each commit is placed in ONE section (first match wins), so order matters.
SECTIONS = OrderedDict([
    ("Booking & Calendar", {
        "keywords": ["booking", "gcal", "google calendar", "calendar import",
                      "calendar integration", "calendar invite",
                      "slot", "availability", "reschedule", "cancel booking",
                      "event type", "event-type", "date exclusion", "freebusy",
                      "office 365", "microsoft webhook", "coachslug",
                      "publicslug", "book a session"],
        "scopes": ["booking"],
        "icon": "",
    }),
    ("Coaching Module", {
        "keywords": ["coaching", "engagement", "coachee", "pre-session",
                      "session card", "coach cancel", "sponsor", "iscoachee",
                      "idp", "grow model", "coachees menu", "coachees list"],
        "scopes": ["coaching", "coach", "journal"],
        "icon": "",
    }),
    ("Conflict Intelligence", {
        "keywords": ["conflict", "risk score", "escalat", "sub-analysis",
                      "sub-analyses", "narrative", "manager script"],
        "scopes": ["conflict"],
        "icon": "",
    }),
    ("AI & Intake Instruments", {
        "keywords": ["ai analysis", "ai usage", "prompt", "hnp", "intake",
                      "instrument", "assessment", "survey template",
                      "custom ai", "claude api", "minresponses"],
        "scopes": ["surveys", "intake"],
        "icon": "",
    }),
    ("Internationalisation (i18n)", {
        "keywords": ["i18n", "translat", "locale", "language switcher",
                      "flag picker", "l10n", "ngx-translate", "i18next",
                      "spanish", "french", "locale-aware", "string extraction",
                      "missing translation"],
        "scopes": ["i18n"],
        "icon": "",
    }),
    ("Mobile App (Ionic / Capacitor)", {
        "keywords": ["mobile", "ionic", "capacitor", "ios", "android",
                      "push notif", "biometric", "haptic", "cocoapod",
                      "firebase", "native"],
        "scopes": ["mobile"],
        "icon": "",
    }),
    ("Authentication & Security", {
        "keywords": ["login page", "login session", "logout", "jwt",
                      "2fa", "recaptcha", "captcha", "cors", "xss", "csp",
                      "security", "rate limit", "trust proxy",
                      "inactivity", "token refresh", "token exchange",
                      "msal", "oauth", "azure client secret"],
        "scopes": ["auth"],
        "icon": "",
    }),
    ("Message Hub & Notifications", {
        "keywords": ["hub", "message hub", "notification pref", "unread",
                      "mark as read", "email mirror", "granular notification",
                      "email communication"],
        "scopes": ["hub"],
        "icon": "",
    }),
    ("Document Generation & Billing", {
        "keywords": ["pdf", "docx", "toolkit", "s3", "download", "invoice",
                      "billing", "stripe", "tax", "company address"],
        "scopes": [],
        "icon": "",
    }),
    ("UI / UX Improvements", {
        "keywords": ["ui(", "ux(", "ux:", "ui:", "carousel", "login page",
                      "split-screen", "hero image", "accordion", "icon button",
                      "mat-icon", "sidebar", "org theme", "org chart",
                      "legal page", "cookie consent", "footer"],
        "scopes": [],
        "icon": "",
    }),
    ("Admin & User Management", {
        "keywords": ["admin", "user management", "role management",
                      "org-settings", "org settings", "permission",
                      "requirepermission", "role-guard"],
        "scopes": ["admin", "users", "org-settings"],
        "icon": "",
    }),
    ("Infrastructure & CI/CD", {
        "keywords": ["ci", "build", "deploy", "docker", "node 22", "npm",
                      "budget", "changelog", "github action", "webpack",
                      "bundle"],
        "scopes": ["ci"],
        "icon": "",
    }),
])


def get_commits(since: str, until: str) -> list[dict]:
    """Fetch commits from git log between since and until dates."""
    cmd = [
        "git", "log",
        "--format=%H%x1f%an%x1f%ad%x1f%s%x1f%b%x1e",
        "--date=short",
        f"--since={since}",
        f"--until={until}",
        "--reverse",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
    commits = []
    for record in result.stdout.strip().split("\x1e"):
        record = record.strip()
        if not record:
            continue
        parts = record.split("\x1f")
        commits.append({
            "hash":    parts[0].strip()[:8] if len(parts) > 0 else "",
            "author":  parts[1].strip() if len(parts) > 1 else "",
            "date":    parts[2].strip() if len(parts) > 2 else "",
            "subject": parts[3].strip() if len(parts) > 3 else "",
            "body":    parts[4].strip() if len(parts) > 4 else "",
        })
    return commits


def parse_conventional(subject: str) -> tuple[str, str, str]:
    """Parse conventional commit: returns (type, scope, description)."""
    m = re.match(
        r"^(feat|fix|chore|docs|refactor|test|style|perf|ci|build|security|ux|ui)"
        r"(?:\(([^)]+)\))?[!]?:\s*(.*)",
        subject, re.IGNORECASE,
    )
    if m:
        return m.group(1).lower(), (m.group(2) or "").lower(), m.group(3)
    return "", "", subject


def classify(commit: dict) -> str:
    """Assign a commit to a single section (first match wins).

    Priority: scope match > subject keywords > body keywords > fallback.
    """
    ctype, scope, desc = parse_conventional(commit["subject"])
    subject_lower = commit["subject"].lower()
    body_lower = commit.get("body", "").lower()

    if commit["subject"].startswith("Merge branch"):
        return "__skip__"

    # 1) Exact scope match (highest priority)
    for section_name, cfg in SECTIONS.items():
        if scope and scope in cfg.get("scopes", []):
            return section_name

    # 2) Keywords matched against subject only (avoids body noise)
    for section_name, cfg in SECTIONS.items():
        if any(kw in subject_lower for kw in cfg["keywords"]):
            return section_name

    # 3) Keywords matched against body (last resort for classification)
    for section_name, cfg in SECTIONS.items():
        if any(kw in body_lower for kw in cfg["keywords"]):
            return section_name

    if ctype in ("fix",):
        return "Bug Fixes & Maintenance"
    if ctype in ("docs", "chore"):
        return "Documentation & Chores"
    return "Other Changes"


def clean_subject(subject: str) -> str:
    """Strip conventional-commit prefix and capitalise."""
    _type, _scope, desc = parse_conventional(subject)
    if desc:
        return desc[0].upper() + desc[1:]
    return subject[0].upper() + subject[1:] if subject else subject


def summarize_body(body: str, max_lines: int = 2) -> str:
    """Extract the first meaningful sentence(s) from commit body."""
    if not body:
        return ""
    lines = []
    for line in body.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("Co-Authored-By:"):
            continue
        if line.startswith("-"):
            line = line.lstrip("- ").strip()
        if len(line) < 10:
            continue
        lines.append(line)
        if len(lines) >= max_lines:
            break
    return " ".join(lines)


def group_commits(commits: list[dict]) -> OrderedDict:
    """Group commits by section, preserving SECTIONS order."""
    groups: dict[str, list[dict]] = {}
    stats = {"feat": 0, "fix": 0, "other": 0}

    for c in commits:
        section = classify(c)
        if section == "__skip__":
            continue
        groups.setdefault(section, []).append(c)
        ctype, _, _ = parse_conventional(c["subject"])
        if ctype == "feat":
            stats["feat"] += 1
        elif ctype == "fix":
            stats["fix"] += 1
        else:
            stats["other"] += 1

    ordered = OrderedDict()
    for name in SECTIONS:
        if name in groups:
            ordered[name] = groups[name]
    for name in ("Bug Fixes & Maintenance", "Documentation & Chores", "Other Changes"):
        if name in groups:
            ordered[name] = groups[name]

    return ordered, stats


def build_doc(groups: OrderedDict, since: str, until: str,
              commit_count: int, stats: dict) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.4)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ARTES Platform — Weekly Changelog")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    title.space_after = Pt(2)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{since}  →  {until}")
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    subtitle.space_after = Pt(4)

    # Stats line
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats_para.add_run(
        f"{commit_count} commits  ·  "
        f"{stats['feat']} features  ·  "
        f"{stats['fix']} fixes  ·  "
        f"{stats['other']} other"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    stats_para.space_after = Pt(12)

    # Sections
    if not groups:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("No changes this period.").font.size = Pt(11)
    else:
        for section_name, commits in groups.items():
            # Section heading
            h = doc.add_paragraph()
            run = h.add_run(section_name)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            h.space_before = Pt(10)
            h.space_after = Pt(3)

            # Add a thin line under the heading
            pf = h.paragraph_format
            from docx.oxml.ns import qn
            pBdr = pf._element.get_or_add_pPr()
            borders = pBdr.makeelement(qn("w:pBdr"), {})
            bottom = borders.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "1",
                qn("w:color"): "3A9FD6",
            })
            borders.append(bottom)
            pBdr.append(borders)

            for c in commits:
                ctype, _, _ = parse_conventional(c["subject"])
                bullet_text = clean_subject(c["subject"])

                # Add type badge
                if ctype == "feat":
                    badge = "NEW"
                elif ctype == "fix":
                    badge = "FIX"
                elif ctype == "refactor":
                    badge = "REFACTOR"
                elif ctype in ("docs", "chore"):
                    badge = "CHORE"
                elif ctype in ("ci", "build"):
                    badge = "CI"
                elif ctype in ("security",):
                    badge = "SECURITY"
                elif ctype in ("ux", "ui"):
                    badge = "UI"
                else:
                    badge = ""

                p = doc.add_paragraph(style="List Bullet")
                p.space_before = Pt(1)
                p.space_after = Pt(1)

                if badge:
                    tag_run = p.add_run(f"[{badge}] ")
                    tag_run.bold = True
                    tag_run.font.size = Pt(8)
                    tag_run.font.color.rgb = BLUE

                main_run = p.add_run(bullet_text)
                main_run.font.size = Pt(9)

                # Show body summary only for features (keeps doc concise)
                if ctype == "feat":
                    summary = summarize_body(c["body"], max_lines=1)
                    if summary and len(summary) > 20:
                        if len(summary) > 150:
                            summary = summary[:147] + "..."
                        detail_para = doc.add_paragraph()
                        detail_para.paragraph_format.left_indent = Inches(0.5)
                        detail_para.space_before = Pt(0)
                        detail_para.space_after = Pt(1)
                        detail_run = detail_para.add_run(summary)
                        detail_run.font.size = Pt(8)
                        detail_run.font.color.rgb = GREY
                        detail_run.italic = True

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.space_before = Pt(16)
    run = footer.add_run(
        "HeadSoft Technology × Helena Coaching  ·  artes.helenacoaching.com"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = LGREY

    return doc


def main():
    parser = argparse.ArgumentParser(description="Generate ARTES weekly changelog")
    parser.add_argument("--days", type=int, default=7,
                        help="How many days back (default: 7). Ignored if --since is set.")
    parser.add_argument("--since", type=str, default=None,
                        help="Start date YYYY-MM-DD (inclusive)")
    parser.add_argument("--until", type=str, default=None,
                        help="End date YYYY-MM-DD (exclusive, default: tomorrow)")
    args = parser.parse_args()

    if args.since:
        since = args.since
    else:
        since = (date.today() - timedelta(days=args.days)).isoformat()

    if args.until:
        until = args.until
    else:
        until = (date.today() + timedelta(days=1)).isoformat()

    commits = get_commits(since, until)
    groups, stats = group_commits(commits)
    doc = build_doc(groups, since, until, len(commits), stats)

    out = f"docs/changelog_{date.today().isoformat()}.docx"
    doc.save(out)
    print(f"Saved: {out}")
    print(f"  {len(commits)} commits across {len(groups)} sections")
    print(f"  {stats['feat']} features, {stats['fix']} fixes, {stats['other']} other")


if __name__ == "__main__":
    main()
