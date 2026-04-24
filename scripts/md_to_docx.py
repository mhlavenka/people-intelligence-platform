"""Focused Markdown-to-DOCX converter for design proposal docs.

Handles the subset of Markdown actually used in docs/design/*.md:
- # / ## / ### / #### headings
- paragraphs with inline **bold**, *italic*, `code`
- unordered lists (- or *)
- numbered lists (1. 2.)
- fenced code blocks (```...```)
- GFM-style tables
- thematic breaks (---)
- block quotes (>)

Usage: python scripts/md_to_docx.py <input.md> <output.docx>
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
ULIST_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
OLIST_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
FENCE_RE = re.compile(r"^```")
HR_RE = re.compile(r"^---+\s*$")


def add_inline(paragraph, text):
    """Parse **bold**, *italic*, `code` into runs."""
    # Tokenize on ** / * / `
    # Simple approach: iterate with a regex that captures the three formats
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\n]+\*)")
    idx = 0
    for m in token_re.finditer(text):
        if m.start() > idx:
            paragraph.add_run(text[idx:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            # shade background
            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            rPr.append(shd)
        elif tok.startswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        idx = m.end()
    if idx < len(text):
        paragraph.add_run(text[idx:])


def add_heading(doc, level, text):
    h = doc.add_heading("", level=level)
    add_inline(h, text)
    return h


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_inline(p, text)
    return p


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    # give it a light grey shading
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc, rows):
    """rows: list of lists of cell strings; first row is header."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            # Clear default paragraph
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], cell_text)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    # Add spacing after table
    doc.add_paragraph()


def parse_table_row(line):
    # trim leading/trailing pipes and split
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Tune default normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    N = len(lines)
    while i < N:
        line = lines[i]

        # Fenced code block
        if FENCE_RE.match(line):
            buf = []
            i += 1
            while i < N and not FENCE_RE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            add_code_block(doc, buf)
            i += 1  # skip closing fence
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            add_heading(doc, level, m.group(2).strip())
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(line):
            p = doc.add_paragraph()
            p_fmt = p.paragraph_format
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "888888")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table: line with pipes and next line is separator
        if "|" in line and i + 1 < N and TABLE_SEP_RE.match(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < N and "|" in lines[i] and lines[i].strip():
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Unordered list
        if ULIST_RE.match(line):
            while i < N and ULIST_RE.match(lines[i]):
                m = ULIST_RE.match(lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, m.group(2).strip())
                i += 1
            continue

        # Ordered list
        if OLIST_RE.match(line):
            while i < N and (OLIST_RE.match(lines[i]) or (lines[i].startswith("   ") and lines[i].strip())):
                m = OLIST_RE.match(lines[i])
                if m:
                    p = doc.add_paragraph(style="List Number")
                    add_inline(p, m.group(2).strip())
                else:
                    # continuation text of previous list item
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    add_inline(p, lines[i].strip())
                i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Block quote
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, line.lstrip("> ").strip())
            i += 1
            continue

        # Paragraph (may span multiple lines)
        buf = [line]
        i += 1
        while i < N and lines[i].strip() and not HEADING_RE.match(lines[i]) and not FENCE_RE.match(lines[i]) \
                and not ULIST_RE.match(lines[i]) and not OLIST_RE.match(lines[i]) and "|" not in lines[i]:
            buf.append(lines[i])
            i += 1
        add_paragraph(doc, " ".join(s.strip() for s in buf))

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
