"""
Rebuild docs/design/technical-howto.docx with current project state.
Preserves the visual style from the original document.
"""
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy, os

SRC = os.path.join(os.path.dirname(__file__), '..', '..', 'docs', 'design', 'technical-howto.docx')
DST = SRC  # overwrite

# Load original to copy styles
orig = Document(SRC)
doc = Document(SRC)

# Clear all content
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

NAVY = RGBColor(0x1B, 0x2A, 0x47)
BLUE = RGBColor(0x3A, 0x9F, 0xD6)
GRAY = RGBColor(0x88, 0x88, 0x88)
DARK = RGBColor(0x22, 0x22, 0x22)
BODY = RGBColor(0x4A, 0x4A, 0x4A)
CODE_BG = "E8F0FE"

def add_title(doc, text, size, color, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return p

def add_body(doc, text, bold=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.font.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def add_code_block(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_info_box(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f'{label}  ')
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.color.rgb = BODY
    return p

def add_list_item(doc, text):
    p = doc.add_paragraph(style='List Paragraph')
    run = p.add_run(text)
    return p

def add_table(doc, rows):
    """rows = list of (col1, col2) tuples"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (c1, c2) in enumerate(rows):
        cell0 = table.cell(i, 0)
        cell1 = table.cell(i, 1)
        r0 = cell0.paragraphs[0].add_run(c1)
        r0.font.bold = True
        r0.font.size = Pt(11)
        r0.font.color.rgb = NAVY
        r1 = cell1.paragraphs[0].add_run(c2)
        r1.font.size = Pt(9)
        r1.font.color.rgb = DARK
    return table

# ═══════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════

# Title block
add_title(doc, 'ARTES — People Intelligence Platform', 28, NAVY, bold=True)
add_title(doc, 'Technical How-To Guide', 16, BLUE)
add_title(doc, 'Setup · Development · Build · Deployment', 12, GRAY)
doc.add_paragraph()

# Meta table
add_table(doc, [
    ('Project', 'ARTES — People Intelligence Platform'),
    ('Client', 'HeadSoft Tech × Helena Coaching'),
    ('Prepared', 'April 2026'),
    ('Version', '2.0'),
    ('Live URL', 'https://artes.helenacoaching.com'),
])
doc.add_paragraph()

# ── 1. System Overview ──
doc.add_heading('1. System Overview', level=1)
add_body(doc, 'ARTES is a multi-tenant B2B SaaS application built for organizational coaching and people analytics. It covers the following modules:')
doc.add_paragraph()
for mod in [
    'Conflict Intelligence — survey-based conflict risk analysis with AI narrative',
    'Neuroinclusion Assessment — maturity gap analysis with AI recommendations',
    'Leadership & Succession — IDP generation using the GROW model',
    'Coaching — engagement tracking, session management, IDP pairing',
    'Booking — coach availability, public booking links, Google/Microsoft Calendar sync',
    'EQ-i Import — psychometric score import and visualization',
    'Journal — session notes, reflective entries for coaches',
    'Survey & Intake — self-service surveys, coach-led interviews, assessments',
    'Message Hub — in-app messaging and notifications',
    'Sponsor Management — external sponsor contacts per engagement',
    'Billing — Stripe-based subscription management',
    'System Administration — cross-org management, invoice generation',
]:
    add_list_item(doc, mod)

doc.add_heading('1.1 Architecture', level=2)
add_table(doc, [
    ('Frontend', 'Angular 17+ (standalone components, signals API)'),
    ('Backend', 'Node.js 20 + Express + TypeScript strict mode'),
    ('Database', 'MongoDB Atlas (mongoose, multi-tenant via organizationId)'),
    ('AI Engine', 'Anthropic Claude API (claude-sonnet-4-6)'),
    ('Auth', 'JWT (15 min access) + refresh tokens (7 days) + optional passkeys (WebAuthn)'),
    ('Email', 'AWS SES via nodemailer'),
    ('Storage', 'AWS S3 (file uploads), base64 for logos'),
    ('Payments', 'Stripe (webhooks at /api/billing/webhook)'),
    ('Calendar', 'Google Calendar API + Microsoft Graph API (bidirectional sync)'),
    ('i18n', 'Frontend: @ngx-translate (en/fr/es) · Backend: i18next'),
    ('Hosting', 'AWS EC2 Amazon Linux + Apache reverse proxy + PM2'),
])

doc.add_heading('1.2 Repository Layout', level=2)
doc.add_paragraph()
add_code_block(doc, """people-intelligence-platform/
├── frontend/
│   ├── src/
│   │   ├── app/
│   │   │   ├── core/              # services, guards, interceptors
│   │   │   ├── shared/            # shared components (avatar, etc.)
│   │   │   └── modules/           # feature modules (see below)
│   │   ├── assets/
│   │   │   └── i18n/              # en.json, fr.json, es.json
│   │   └── environments/
│   ├── angular.json
│   └── package.json
├── backend/
│   ├── src/
│   │   ├── config/                # env.ts, database.ts
│   │   ├── controllers/
│   │   ├── middleware/            # auth, tenantResolver, error, rateLimiter
│   │   ├── models/                # Mongoose models (28+)
│   │   ├── routes/                # one file per domain (27 route files)
│   │   ├── services/              # ai, email, booking, calendar, etc.
│   │   ├── jobs/                  # cron jobs (reminders, webhooks, trials)
│   │   ├── scripts/               # seed & migration scripts
│   │   └── locales/               # i18next locale files (en/fr/es)
│   ├── ecosystem.config.js        # PM2 config
│   ├── tsconfig.json
│   ├── tsconfig.build.json
│   └── package.json
├── docs/                          # design docs, prompts, data feeds
├── scripts/                       # utility scripts
├── deploy.sh                      # automated deploy script
└── docker-compose.yml             # optional local Docker stack""")

doc.add_paragraph()
add_body(doc, 'Frontend modules:', bold=True)
add_body(doc, 'admin, auth, billing, booking, coach, coaching, conflict, dashboard, eq-import, hub, journal, legal, neuroinclusion, org-chart, profile, settings, sponsor, succession, survey, system-admin')

# ── 2. Prerequisites ──
doc.add_heading('2. Prerequisites', level=1)
doc.add_heading('2.1 Local Machine', level=2)
add_table(doc, [
    ('Node.js', '≥ 20.x  (check: node --version)'),
    ('npm', '≥ 10.x  (check: npm --version)'),
    ('Angular CLI', 'npm install -g @angular/cli'),
    ('Git', '≥ 2.x'),
    ('WebStorm / VS Code', 'WebStorm 2023.3+ or VS Code with Angular extension'),
    ('MongoDB', 'Atlas account or local mongod for dev'),
    ('Python', '≥ 3.10 (for docx scripts, optional)'),
])

doc.add_heading('2.2 Server Requirements', level=2)
add_table(doc, [
    ('OS', 'Amazon Linux 2023'),
    ('Web server', 'Apache 2.4 (httpd) with mod_proxy'),
    ('Node', '≥ 20.x (installed via nvm)'),
    ('PM2', 'npm install -g pm2'),
    ('PEM key', 'headsoft-aws.pem (stored locally, never committed)'),
])

doc.add_heading('2.3 Environment Variables', level=2)
add_body(doc, 'Create a .env file in the backend/ folder (never commit this file):')
doc.add_paragraph()
add_code_block(doc, """# backend/.env
NODE_ENV=development
PORT=3030
FRONTEND_URL=http://localhost:4200
API_BASE_URL=http://localhost:3030
PUBLIC_API_BASE_URL=http://localhost:3030

MONGODB_URI=mongodb+srv://<user>:<pass>@cluster.mongodb.net/artes

JWT_SECRET=<long-random-string>
JWT_EXPIRES_IN=15m
JWT_REFRESH_SECRET=<another-long-random-string>
JWT_REFRESH_EXPIRES_IN=7d
CANCEL_TOKEN_JWT_SECRET=<random-string>

ANTHROPIC_API_KEY=sk-ant-...

AWS_REGION=us-east-1
AWS_ACCESS_KEY_ID=...
AWS_SECRET_ACCESS_KEY=...
AWS_S3_BUCKET=artes-uploads
AWS_SES_FROM_EMAIL=noreply@helenacoaching.com

STRIPE_SECRET_KEY=sk_live_...
STRIPE_WEBHOOK_SECRET=whsec_...

# Google Calendar integration
GOOGLE_CLIENT_ID=...
GOOGLE_CLIENT_SECRET=...
GOOGLE_CALENDAR_REDIRECT_URI=http://localhost:3030/api/calendar/google/callback
GOOGLE_WEBHOOK_SECRET=<random-string>

# Microsoft Calendar integration
MICROSOFT_CLIENT_ID=...
MICROSOFT_CLIENT_SECRET=...
MICROSOFT_TENANT_ID=...
MICROSOFT_CALENDAR_REDIRECT_URI=http://localhost:3030/api/calendar/microsoft/callback
MICROSOFT_WEBHOOK_SECRET=<random-string>

# Passkey / WebAuthn
WEBAUTHN_RP_NAME=ARTES
WEBAUTHN_RP_ID=localhost
WEBAUTHN_ORIGIN=http://localhost:4200

# Booking webhooks (enable on production after proxy setup)
BOOKING_WEBHOOKS_ENABLED=false""")

doc.add_paragraph()
add_info_box(doc, '⚠ Note:', 'backend/.env is already listed in .gitignore. Never commit real secrets.')

# ── 3. Running Locally ──
doc.add_heading('3. Running Locally', level=1)
doc.add_heading('3.1 Clone & Install', level=2)
doc.add_paragraph()
add_code_block(doc, """git clone https://github.com/mhlavenka/people-intelligence-platform.git
cd people-intelligence-platform

# Install backend dependencies
cd backend && npm install

# Install frontend dependencies
cd ../frontend && npm install""")

doc.add_heading('3.2 Start the Backend', level=2)
doc.add_paragraph()
add_code_block(doc, """cd backend
npm run dev
# → ts-node-dev watches src/ and restarts on changes
# → API available at http://localhost:3030/api""")

doc.add_heading('3.3 Start the Frontend', level=2)
doc.add_paragraph()
add_code_block(doc, """cd frontend
npx ng serve
# → App available at http://localhost:4200
# → Proxy: /api → http://localhost:3030/api  (see proxy.conf.json)""")
doc.add_paragraph()
add_info_box(doc, 'ℹ Info:', 'The Angular dev server proxies /api calls to the backend automatically via proxy.conf.json. You do not need to change any URLs.')

doc.add_heading('3.4 Seed Initial Data', level=2)
add_body(doc, 'Run once after setting up a fresh database:')
doc.add_paragraph()
add_code_block(doc, """# Create the first super-admin user + organization
cd backend && npm run seed:admin

# Seed global survey/intake templates
npm run seed:surveys

# Seed subscription plans (12 plans)
npm run seed:plans""")
doc.add_paragraph()
add_body(doc, 'On the server (compiled JS only — no ts-node):')
add_code_block(doc, """node dist/scripts/seed-admin.js
node dist/scripts/seed-plans.js
node dist/scripts/seed-surveys.js""")

doc.add_heading('3.5 Opening in WebStorm', level=2)
add_list_item(doc, 'Open the repo root folder in WebStorm (File → Open).')
add_list_item(doc, 'WebStorm will detect both package.json files automatically.')
add_list_item(doc, 'Go to Run → Edit Configurations → + → npm:')
add_list_item(doc, 'Name: Backend Dev | package.json: backend/package.json | Script: dev')
add_list_item(doc, 'Name: Frontend Dev | package.json: frontend/package.json | Script: start (or use ng serve)')
add_list_item(doc, 'Run both configs simultaneously using the Services panel (View → Tool Windows → Services).')
add_list_item(doc, 'Enable TypeScript service: Settings → Languages → TypeScript → Use TypeScript from node_modules.')
add_list_item(doc, 'For Angular: install the Angular and Angular CLI plugins from the WebStorm Marketplace.')

# ── 4. Building for Production ──
doc.add_heading('4. Building for Production', level=1)
doc.add_heading('4.1 Frontend Build', level=2)
doc.add_paragraph()
add_code_block(doc, """cd frontend
npx ng build --configuration production
# Output: frontend/dist/artes-frontend/browser/""")
doc.add_paragraph()
add_body(doc, 'The production build uses src/environments/environment.prod.ts which sets apiUrl to the relative path /api — this routes through Apache\'s reverse proxy on the server.')

doc.add_heading('4.2 Backend Build', level=2)
add_body(doc, 'The backend uses a separate tsconfig.build.json that disables declaration files and source maps for faster compilation:')
doc.add_paragraph()
add_code_block(doc, """cd backend
npm run build
# → tsc -p tsconfig.build.json
# Output: backend/dist/""")
doc.add_paragraph()
add_info_box(doc, 'ℹ Info:', 'tsconfig.build.json extends the base tsconfig but sets declaration: false, declarationMap: false, sourceMap: false. This dramatically speeds up builds on low-RAM servers (e.g. t2.micro EC2).')

# ── 5. Deployment to EC2 ──
doc.add_heading('5. Deployment to EC2', level=1)
doc.add_heading('5.1 Automated Deploy Script', level=2)
add_body(doc, 'The repository includes deploy.sh in the project root. It handles the entire build + upload + restart cycle in one command.')
doc.add_paragraph()
add_code_block(doc, """# Full deploy: frontend + backend
bash deploy.sh all

# Frontend only (CSS / template changes)
bash deploy.sh frontend

# Backend only (API / model changes)
bash deploy.sh backend""")
doc.add_paragraph()
add_body(doc, 'The script performs these steps:')
add_list_item(doc, 'Checks translation file sync (en/fr/es key counts must match)')
add_list_item(doc, 'Frontend: ng build --configuration production → scp browser/ to /opt/apps/artes/frontend/public/')
add_list_item(doc, 'Backend: npm run build (local tsc) → scp dist/ + package.json to server → npm install --omit=dev → pm2 restart artes-backend')

doc.add_heading('5.2 Manual Deploy (step by step)', level=2)
doc.add_heading('5.2.1 Upload Frontend', level=3)
doc.add_paragraph()
add_code_block(doc, """cd frontend
npx ng build --configuration production
scp -i headsoft-aws.pem -r dist/artes-frontend/browser/. \\
  ec2-user@13.218.6.173:/opt/apps/artes/frontend/public/""")

doc.add_heading('5.2.2 Upload & Restart Backend', level=3)
doc.add_paragraph()
add_code_block(doc, """cd backend
npm run build
scp -i headsoft-aws.pem -r dist/. ec2-user@13.218.6.173:/opt/apps/artes/backend/dist/
scp -i headsoft-aws.pem package.json package-lock.json ec2-user@13.218.6.173:/opt/apps/artes/backend/
ssh -i headsoft-aws.pem ec2-user@13.218.6.173 \\
  "cd /opt/apps/artes/backend && npm install --omit=dev && pm2 restart artes-backend && pm2 save\"""")

doc.add_heading('5.3 PM2 Process Management', level=2)
add_body(doc, 'The backend runs under PM2 using ecosystem.config.js:')
doc.add_paragraph()
add_code_block(doc, """# Check status
pm2 status

# View logs
pm2 logs artes-backend
pm2 logs artes-backend --lines 100

# Restart / reload
pm2 restart artes-backend
pm2 reload artes-backend   # zero-downtime reload

# Start for the first time (on the server)
cd /opt/apps/artes/backend
pm2 start ecosystem.config.js --env production
pm2 save
pm2 startup              # enable auto-start on reboot""")

doc.add_heading('5.4 Apache Configuration', level=2)
add_body(doc, 'Apache acts as a reverse proxy: it serves the Angular SPA and forwards /api/* to Node.js on port 3030.')
doc.add_paragraph()
add_code_block(doc, """# /etc/httpd/conf.d/artes.conf
<VirtualHost *:443>
  ServerName artes.helenacoaching.com

  DocumentRoot /opt/apps/artes/frontend/public

  # Angular SPA routing
  <Directory "/opt/apps/artes/frontend/public">
    Options -Indexes
    AllowOverride None
    Require all granted
    FallbackResource /index.html
  </Directory>

  # Backend API proxy
  ProxyPreserveHost On
  ProxyPass /api http://localhost:3030/api
  ProxyPassReverse /api http://localhost:3030/api

  # SSL (managed by Certbot)
  SSLEngine on
  SSLCertificateFile    /etc/letsencrypt/live/artes.helenacoaching.com/fullchain.pem
  SSLCertificateKeyFile /etc/letsencrypt/live/artes.helenacoaching.com/privkey.pem
</VirtualHost>

# Reload Apache after config changes
sudo systemctl reload httpd""")

# ── 6. First-time Server Setup ──
doc.add_heading('6. First-time Server Setup', level=1)
add_body(doc, 'Run these commands once on a fresh EC2 Amazon Linux instance:')

doc.add_heading('6.1 Install Node.js via nvm', level=2)
doc.add_paragraph()
add_code_block(doc, """curl -o- https://raw.githubusercontent.com/nvm-sh/nvm/v0.39.7/install.sh | bash
source ~/.bashrc
nvm install 20
nvm use 20
node --version   # should print v20.x.x""")

doc.add_heading('6.2 Install PM2 and Apache', level=2)
doc.add_paragraph()
add_code_block(doc, """npm install -g pm2
sudo yum install -y httpd mod_ssl
sudo systemctl enable httpd
sudo systemctl start httpd""")

doc.add_heading('6.3 Enable Apache Proxy Modules', level=2)
doc.add_paragraph()
add_code_block(doc, """sudo yum install -y mod_proxy mod_proxy_http
# These are usually already enabled on Amazon Linux""")

doc.add_heading('6.4 Create App Directories', level=2)
doc.add_paragraph()
add_code_block(doc, """sudo mkdir -p /opt/apps/artes/frontend/public
sudo mkdir -p /opt/apps/artes/backend
sudo chown -R ec2-user:ec2-user /opt/apps/artes""")

doc.add_heading('6.5 Copy .env to Server', level=2)
doc.add_paragraph()
add_code_block(doc, """# Run from your local machine
scp -i headsoft-aws.pem backend/.env ec2-user@13.218.6.173:/opt/apps/artes/backend/.env""")

doc.add_heading('6.6 Initial Backend Start', level=2)
doc.add_paragraph()
add_code_block(doc, """# After first deploy (bash deploy.sh backend), start PM2:
cd /opt/apps/artes/backend
pm2 start ecosystem.config.js --env production
pm2 save
pm2 startup   # follow the printed command to enable auto-start""")

# ── 7. Environment Configuration ──
doc.add_heading('7. Environment Configuration', level=1)
doc.add_heading('7.1 Frontend Environments', level=2)
add_body(doc, 'Angular automatically selects the correct environment file based on the build configuration:')
doc.add_paragraph()
add_table(doc, [
    ('Development', 'src/environments/environment.ts → apiUrl: "http://localhost:3030/api"'),
    ('Production', 'src/environments/environment.prod.ts → apiUrl: "/api" (relative, via Apache proxy)'),
])
doc.add_paragraph()
add_code_block(doc, """npx ng serve                            # uses environment.ts
npx ng build --configuration production # uses environment.prod.ts""")

doc.add_heading('7.2 Backend Config (backend/src/config/env.ts)', level=2)
add_body(doc, 'All configuration is read from environment variables with sensible defaults:')
doc.add_paragraph()
add_table(doc, [
    ('PORT', '3030'),
    ('NODE_ENV', 'development'),
    ('JWT_EXPIRES_IN', '15m'),
    ('JWT_REFRESH_EXPIRES_IN', '7d'),
    ('AWS_REGION', 'us-east-1'),
    ('BOOKING_WEBHOOKS_ENABLED', 'false'),
])

# ── 8. Key Architectural Decisions ──
doc.add_heading('8. Key Architectural Decisions', level=1)

doc.add_heading('8.1 Multi-tenancy', level=2)
add_body(doc, 'Every MongoDB document (except global templates and admin-level entities) carries an organizationId field. The tenantFilter mongoose plugin warns at runtime if a query is issued without this filter. All API routes go through the auth middleware which injects req.user.organizationId from the JWT.')

doc.add_heading('8.2 JWT Auth Flow', level=2)
add_list_item(doc, 'Access token: 15 minute lifetime, sent in Authorization: Bearer <token> header.')
add_list_item(doc, 'Refresh token: 7 day lifetime, stored in the database, used to obtain a new access token.')
add_list_item(doc, 'Angular authInterceptor automatically intercepts 401 responses, calls /auth/refresh, and retries the original request.')
add_list_item(doc, 'Proactive refresh: AuthService.scheduleTokenRefresh() refreshes the token 60 seconds before expiry.')
add_list_item(doc, 'Inactivity logout: 30 minutes of inactivity with a 2-minute warning dialog.')

doc.add_heading('8.3 Passkey / WebAuthn Authentication', level=2)
add_body(doc, 'Users can register passkeys (FIDO2/WebAuthn) as a second factor or passwordless login method. Registration and authentication endpoints live under /api/auth/passkey/*. Credentials are stored per-user in the database.')

doc.add_heading('8.4 Survey Privacy — Minimum Group Size', level=2)
add_body(doc, 'Survey response aggregation enforces a minimum group size of 5 respondents. Raw individual responses are never returned to prevent de-anonymisation.')

doc.add_heading('8.5 Global Survey Templates', level=2)
add_body(doc, 'Survey templates can be either org-scoped (organizationId set) or global (isGlobal: true, no organizationId). The survey route queries both: { $or: [{ organizationId }, { isGlobal: true }] }. Global templates are seeded via npm run seed:surveys.')

doc.add_heading('8.6 Organization Logo Storage', level=2)
add_body(doc, 'Logos are stored as base64 Data URLs directly in MongoDB (logoUrl field on the Organization model). The client validates a 2 MB maximum before upload. No S3 bucket is used for logos.')

doc.add_heading('8.7 Internationalization (i18n)', level=2)
add_list_item(doc, 'Frontend: All UI strings use {{ \'KEY\' | translate }} — never hardcode English in Angular templates.')
add_list_item(doc, 'Translation files: frontend/src/assets/i18n/en.json, fr.json, es.json (2400+ keys each).')
add_list_item(doc, 'Backend: req.t(\'errors.xxx\') via i18next; locale files in backend/src/locales/{lang}/common.json.')
add_list_item(doc, 'AI prompts: all prompt builders accept a language parameter — always pass req.language from the route handler.')
add_list_item(doc, 'Run node scripts/i18n/check-translations.js before deploying to catch missing keys.')

doc.add_heading('8.8 Roles & Access', level=2)
add_body(doc, 'The application defines the following roles:')
add_table(doc, [
    ('system_admin', 'Cross-org system administration only (separate shell at /system-admin)'),
    ('admin', 'Full org access'),
    ('hr_manager', 'All operational modules, no billing/org settings'),
    ('manager', 'Read access to conflict/neuro-inclusion analytics, can escalate'),
    ('coach', 'Conduct interviews, view intake templates, succession/IDP, booking, coaching'),
    ('coachee', 'Take surveys, view own IDP (coachee is also a flag: User.isCoachee, orthogonal to role)'),
])

doc.add_heading('8.9 Booking & Calendar Integration', level=2)
add_body(doc, 'The booking module provides coach availability management and public booking links with Google Calendar and Microsoft Calendar bidirectional sync.')
add_list_item(doc, 'Coaches configure availability windows, exclusion dates, and buffer times via booking settings.')
add_list_item(doc, 'Public booking page: /book/:coachSlug — anonymous or authenticated booking.')
add_list_item(doc, 'Google Calendar events are created/updated/deleted in sync with bookings.')
add_list_item(doc, 'GCal push notifications (webhooks) detect coach-side changes (cancellation, reschedule).')
add_list_item(doc, 'Microsoft Calendar integration follows the same pattern via Graph API subscriptions.')
add_list_item(doc, 'Webhook endpoints: POST /api/webhooks/gcal and POST /api/webhooks/microsoft.')
add_list_item(doc, 'Enable with BOOKING_WEBHOOKS_ENABLED=true after the Apache vhost proxies the webhook paths to PM2.')

doc.add_heading('8.10 Coaching Module', level=2)
add_body(doc, 'Coaching engagements track coach-coachee relationships with sessions, IDPs, and optional sponsor contacts.')
add_list_item(doc, 'CoachingEngagement: links coach + coachee + organization with status tracking.')
add_list_item(doc, 'CoachingSession: individual sessions with date, duration, notes, status (scheduled/completed/cancelled).')
add_list_item(doc, 'Booking ↔ Session sync: bookings with an internal coachee auto-create paired CoachingSession records.')
add_list_item(doc, 'Cancel/reschedule propagation works bidirectionally between Booking and CoachingSession.')

doc.add_heading('8.11 Background Jobs', level=2)
add_body(doc, 'Cron jobs run inside the backend process (no separate worker):')
add_table(doc, [
    ('reminder.job.ts', 'Booking reminders — 24h and 1h before session (every 15 min)'),
    ('webhookRenewal.job.ts', 'Renews Google/Microsoft push notification channels before expiry (hourly)'),
    ('preSessionIntake.job.ts', 'Sends pre-session intake surveys to coachees before upcoming sessions'),
    ('trialRevert.job.ts', 'Reverts expired trial subscriptions back to free plan'),
])

# ── 9. Database Models ──
doc.add_heading('9. Database Models', level=1)
add_body(doc, 'Key Mongoose models (28+):')
add_table(doc, [
    ('Organization', 'name, slug, billingEmail, industry, employeeCount, plan, departments[]'),
    ('User', 'organizationId, email, passwordHash, role, department, isCoachee, twoFactorEnabled, preferredLanguage'),
    ('SurveyTemplate', 'organizationId, title, moduleType, intakeType (survey|interview|assessment), questions[], isActive'),
    ('SurveyResponse', 'organizationId, templateId, respondentId?, coachId?, sessionFormat, targetName, submissionToken'),
    ('ConflictAnalysis', 'organizationId, riskScore, riskLevel, conflictTypes[], aiNarrative, managerScript'),
    ('NeuroinclustionAssessment', 'organizationId, individual/org scores, gapAnalysis'),
    ('DevelopmentPlan', 'organizationId, userId, goals[], grow model fields'),
    ('CoachingEngagement', 'organizationId, coachId, coacheeId, status, sessionCount'),
    ('CoachingSession', 'engagementId, date, duration, status, bookingId, notes'),
    ('Booking', 'organizationId, coachId, coacheeId?, startTime, endTime, googleEventId, status'),
    ('BookingSettings', 'coachId, availability, bufferMinutes, targetCalendarId, conflictCalendarIds'),
    ('AvailabilityConfig', 'coachId, dayOfWeek, startTime, endTime, exclusionDates'),
    ('Message', 'organizationId, fromUserId, toUserId, content, isRead'),
    ('Notification', 'organizationId, userId, type, title, body, isRead, link'),
    ('JournalEntry / JournalSessionNote', 'organizationId, coachId, content, tags'),
    ('Plan', 'key (unique), name, priceMonthly (cents), maxUsers, features[], isActive'),
    ('Invoice', 'organizationId, amount, status, planKey'),
    ('Sponsor', 'organizationId, engagementId, firstName, lastName, email, company'),
    ('EqiScoreRecord', 'organizationId, userId, scores, importedAt'),
    ('WebhookState', 'coachId, provider, channelId, resourceId, expiration'),
])

# ── 10. API Routes Summary ──
doc.add_heading('10. API Routes Summary', level=1)
doc.add_paragraph()
add_code_block(doc, """/api/auth/*                    — login, register, refresh, 2FA, passkeys
/api/auth/oauth/*              — Google/Microsoft OAuth flows
/api/users/*                   — CRUD (admin/hr), /me (self), /coachees (coach+)
/api/surveys/*                 — templates CRUD, /respond, /check/:id, /responses/:id
/api/conflict/*                — /analyze, /analyses, /analyses/:id/sub-analyses, /escalate/:id
/api/neuroinclusion/*          — assessment CRUD + AI gap analysis
/api/succession/*              — IDP CRUD + AI generation
/api/coaching/*                — engagements, sessions, IDP links
/api/booking/*                 — settings, availability, public booking
/api/calendar/*                — Google/Microsoft Calendar connect/disconnect/callback
/api/hub/*                     — messages, notifications, mark read/unread
/api/journal/*                 — entries, session notes, reflective entries
/api/sponsor/*                 — sponsor CRUD per engagement
/api/eqi-import/*              — EQ-i score import + parsing
/api/org-chart/*               — org chart data
/api/dashboard/*               — summary stats
/api/billing/*                 — Stripe integration (org-level)
/api/plans/*                   — plan catalog (GET public, admin CRUD)
/api/reports/*                 — report generation
/api/roles/*                   — custom role management
/api/system-admin/*            — cross-org admin
/api/system-admin/billing/*    — invoice generation
/api/system-admin/settings/*   — global app settings
/api/webhooks/gcal             — Google Calendar push notifications
/api/webhooks/microsoft        — Microsoft Calendar push notifications
/api/public/booking/:slug      — public booking form (no auth required)""")

# ── 11. NPM Scripts Reference ──
doc.add_heading('11. NPM Scripts Reference', level=1)
doc.add_heading('11.1 Backend (backend/)', level=2)
add_table(doc, [
    ('npm run dev', 'Start ts-node-dev with hot reload (development)'),
    ('npm run build', 'Compile TypeScript with tsconfig.build.json'),
    ('npm start', 'Run compiled dist/app.js (production)'),
    ('npm test', 'Run integration tests'),
    ('npm run seed:admin', 'Create first super-admin user + organization in MongoDB'),
    ('npm run seed:surveys', 'Seed global survey/intake templates'),
    ('npm run seed:plans', 'Seed 12 subscription plans (drops existing)'),
    ('npm run lint', 'Run ESLint on src/'),
    ('npm run deploy', 'Build + pm2 restart (run from server)'),
    ('npm run backfill:isCoachee', 'Backfill User.isCoachee flag from existing engagements'),
])

doc.add_heading('11.2 Frontend (frontend/)', level=2)
add_table(doc, [
    ('npx ng serve', 'Start dev server on :4200 with proxy'),
    ('npx ng build --configuration production', 'Production build → dist/artes-frontend/browser/'),
    ('npx ng lint', 'Run ESLint'),
    ('npx ng test', 'Run unit tests (Karma)'),
])

# ── 12. Troubleshooting ──
doc.add_heading('12. Troubleshooting', level=1)

doc.add_heading('12.1 Login calls localhost instead of the real API', level=2)
add_body(doc, 'Cause: Frontend was built without --configuration production, so environment.ts (dev) was used.')
add_body(doc, 'Fix:')
add_code_block(doc, """cd frontend
npx ng build --configuration production
bash deploy.sh frontend""")

doc.add_heading('12.2 PM2 picks bun interpreter', level=2)
add_body(doc, 'Cause: pm2 start src/app.ts — PM2 detects .ts extension and tries to use bun.')
add_body(doc, 'Fix: Always start via ecosystem.config.js which explicitly sets interpreter: "node" and script: "dist/app.js".')
add_code_block(doc, """pm2 start ecosystem.config.js --env production""")

doc.add_heading('12.3 tsc hangs or runs out of memory on EC2', level=2)
add_body(doc, 'Cause: Full tsconfig.json with declaration: true + declarationMap: true generates extra files, slow on t2.micro.')
add_body(doc, 'Fix: Always use npm run build which invokes tsc -p tsconfig.build.json (declaration and sourcemaps disabled).')

doc.add_heading('12.4 Survey templates not visible', level=2)
add_body(doc, 'Cause: Templates were seeded with an organizationId, making them invisible to other orgs.')
add_body(doc, 'Fix: Re-run npm run seed:surveys to recreate them with isGlobal: true.')

doc.add_heading('12.5 CORS errors in development', level=2)
add_body(doc, 'Cause: Direct API calls bypassing the Angular dev proxy.')
add_body(doc, 'Fix: Ensure frontend/proxy.conf.json is configured and ng serve is started with --proxy-config proxy.conf.json (this is already wired in angular.json).')

doc.add_heading('12.6 MongoDB connection refused', level=2)
add_body(doc, 'Verify MONGODB_URI in backend/.env is correct and your IP is whitelisted in MongoDB Atlas (Network Access → Add IP Address).')

doc.add_heading('12.7 Translation keys showing as raw keys in UI', level=2)
add_body(doc, 'Cause: Key was added to one language file but not all three (en.json, fr.json, es.json).')
add_body(doc, 'Fix: Run node scripts/i18n/check-translations.js to find missing keys. The deploy script runs this automatically.')

doc.add_heading('12.8 Google Calendar webhook not receiving events', level=2)
add_body(doc, 'Cause: BOOKING_WEBHOOKS_ENABLED is false, or the Apache vhost is not proxying /api/webhooks/gcal to PM2.')
add_body(doc, 'Fix: Set BOOKING_WEBHOOKS_ENABLED=true in .env and ensure the Apache proxy covers the /api path including webhooks.')

doc.add_paragraph()
doc.add_paragraph()
add_body(doc, 'ARTES — People Intelligence Platform © 2026 HeadSoft Tech. Internal use only.', size=9, color=GRAY)

# Save
doc.save(DST)
print('✓ Document saved to', DST)
